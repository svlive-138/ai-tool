"""
AI Grammar Correction Streamlit app

Usage (PowerShell):
    cd "c:\\shrey\\documents\\my projects\\ai tool"
    streamlit run "spelling_correction_app.py"

Notes:
 - This app uses a transformer model from Hugging Face; ensure you have enough RAM/CPU.
 - NLTK punkt is downloaded at runtime (first run will fetch it).
 - For Windows TTS, `pyttsx3` is used as a server-side fallback; the app prefers browser TTS when available.
"""

import json
import io
import streamlit as st
import streamlit.components.v1 as components
try:
    from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
    TRANSFORMERS_AVAILABLE = True
    _transformers_import_error = None
except Exception as _e:
    TRANSFORMERS_AVAILABLE = False
    _transformers_import_error = _e
import torch
from docx import Document
import nltk
nltk.download('punkt', quiet=True)
from nltk.tokenize import sent_tokenize

if not TRANSFORMERS_AVAILABLE:
    st.error("Transformers library not available. Please install it to enable grammar correction. Run:\n   pip install -U transformers accelerate torch")
    # Provide a small no-op stubs so rest of the app doesn't crash
    def correct_long_text(text: str) -> str:
        return """Transformers not installed. Install transformers and restart the app to enable correction."""
else:
    # Load model and tokenizer (cached to avoid re-loading on every Streamlit rerun)
    @st.cache_resource
    def load_model_and_tokenizer(model_name: str = "vennify/t5-base-grammar-correction"):
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        try:
            model = AutoModelForSeq2SeqLM.from_pretrained(model_name, device_map=None, low_cpu_mem_usage=False)
        except TypeError:
            model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
        return tokenizer, model

    tokenizer, model = load_model_and_tokenizer()

    # Choose runtime device; we'll try to move model but handle meta-tensor errors gracefully
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    try:
        model.to(device)
    except NotImplementedError as e:
        st.warning("Could not move model to the selected device (GPU). Falling back to CPU. If you want GPU inference, try installing a recent 'transformers' and 'accelerate' and load with device_map='auto'.\nError: {}".format(e))
        device = torch.device("cpu")

# (Using browser TTS only; server-side TTS removed for smoother deployment)

# --- core text functions ---

def correct_text(text: str) -> str:
    # legacy compatibility: wrapper kept but prefers explicit params in correct_long_text
    return correct_text_with_params(text, gen_beams=3, max_length=800)


def correct_text_with_params(text: str, gen_beams: int = 3, max_length: int = 800) -> str:
    input_text = "grammar: " + text
    inputs = tokenizer.encode(input_text, return_tensors="pt", max_length=max_length, truncation=True).to(device)
    outputs = model.generate(inputs, max_length=max_length, num_beams=gen_beams, no_repeat_ngram_size=3,
                             length_penalty=1.0, early_stopping=True)
    return tokenizer.decode(outputs[0], skip_special_tokens=True)


def chunk_text(text: str, max_tokens: int = 500):
    paragraphs = [p for p in text.split("\n\n")]
    chunks = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            chunks.append("")
            continue
        sentences = sent_tokenize(para)
        current_chunk = ""
        for sentence in sentences:
            combined = (current_chunk + ' ' + sentence).strip()
            if len(tokenizer.encode(combined)) < max_tokens:
                current_chunk = combined
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
        if current_chunk:
            chunks.append(current_chunk.strip())
    return chunks


def correct_long_text(text: str, gen_beams: int = 3, max_length: int = 800) -> str:
    paragraphs = [p for p in text.split("\n\n")]
    corrected_paras = []
    total_paragraphs = max(1, len(paragraphs))
    progress = st.progress(0)

    for idx, para in enumerate(paragraphs):
        if not para.strip():
            corrected_paras.append("")
            progress.progress((idx + 1) / total_paragraphs)
            continue

        para_chunks = chunk_text(para, max_tokens=500)
        corrected_chunks = []
        for chunk in para_chunks:
            if not chunk:
                continue
            corrected_chunks.append(correct_text_with_params(chunk, gen_beams=gen_beams, max_length=max_length))

        corrected_para = " ".join(corrected_chunks).strip()
        corrected_paras.append(corrected_para)
        progress.progress((idx + 1) / total_paragraphs)

    return "\n\n".join(corrected_paras)


# --- TTS: use client-side Web Speech API (preferred) ---
def speak_in_browser(text: str):
    """Inject Web Speech API JS to the client to speak the provided text.
    Note: browsers may block autoplay; the user can interact with the page to allow playback.
    """
    try:
        safe_text = json.dumps(text)
        js = f"""
        <script>
        const txt = {safe_text};
        const u = new SpeechSynthesisUtterance(txt);
        u.rate = 1.0;
        window.speechSynthesis.cancel();
        window.speechSynthesis.speak(u);
        </script>
        """
        components.html(js, height=0)
    except Exception as e:
        st.warning(f"Browser TTS failed: {e}")


# --- file utilities ---
def process_file(uploaded_file):
    if uploaded_file.type == "text/plain":
        return uploaded_file.read().decode("utf-8")
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        st.error("Unsupported file type. Please upload a .txt or .docx file.")
        return None


def create_docx(text: str):
    doc = Document()
    for para in text.split("\n\n"):
        doc.add_paragraph(para)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- UI ---
st.set_page_config(page_title="AI Grammar Correction", layout="centered")
st.title("📚 AI Grammar Correction Tool")

st.markdown(
    "<style>\n    .stApp { max-width: 1100px; margin: auto; }\n    .title { font-size: 32px; font-weight: 700; }\n    .box { padding: 12px; border-radius: 8px; background: #f7f9fc; }\n    pre { white-space: pre-wrap; font-family: inherit; color: #111; }\n</style>",
    unsafe_allow_html=True,
)

input_mode = st.radio("Choose input method:", ["Text Input", "Upload File"]) 

# Speed/quality selector to trade off latency vs output quality
speed_choice = st.sidebar.selectbox("Response speed", ["Fast", "Balanced", "Quality"], index=1)
if speed_choice == "Fast":
    GEN_BEAMS = 1
    MAX_LENGTH = 400
elif speed_choice == "Quality":
    GEN_BEAMS = 6
    MAX_LENGTH = 1050
else:  # Balanced
    GEN_BEAMS = 3
    MAX_LENGTH = 800

# Browser TTS helper (fast, client-side playback)
def speak_in_browser(text: str):
    try:
        safe_text = json.dumps(text)
        js = f"""
        <script>
        const txt = {safe_text};
        const u = new SpeechSynthesisUtterance(txt);
        u.rate = 1.0;
        window.speechSynthesis.cancel();
        window.speechSynthesis.speak(u);
        </script>
        """
        components.html(js, height=0)
    except Exception as e:
        st.warning(f"Browser TTS failed: {e}")
if input_mode == "Text Input":
    user_input = st.text_area("Enter your text here:")
    if st.button("Correct Grammar"):
        if user_input.strip():
            with st.spinner("Correcting..."):
                corrected = correct_long_text(user_input, gen_beams=6, max_length=5000)

            st.subheader("✅ Result")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**Original**")
                st.text_area("Original text", value=user_input, height=300, key="orig_text", label_visibility="hidden")
            with col2:
                st.markdown("**Corrected**")
                st.text_area("Corrected text", value=corrected, height=300, key="corr_text", label_visibility="hidden")

            # Prefer client-side playback for speed
            try:
                speak_in_browser(corrected)
            except Exception:
                st.info("Browser TTS not available — interact with the page or use a browser with Web Speech API support.")

            docx_file = create_docx(corrected)
            txt_bytes = corrected.encode("utf-8")
            st.download_button(label="📥 Download DOCX", data=docx_file, file_name="corrected_output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.download_button(label="📄 Download TXT", data=txt_bytes, file_name="corrected_output.txt", mime="text/plain")
        else:
            st.warning("Please enter some text.")

else:
    uploaded_file = st.file_uploader("Upload a .txt or .docx file", type=["txt", "docx"])
    if uploaded_file and st.button("Correct File"):
        with st.spinner("Processing file..."):
            raw_text = process_file(uploaded_file)
            if raw_text:
                corrected = correct_long_text(raw_text, gen_beams=GEN_BEAMS, max_length=MAX_LENGTH)

                st.subheader("✅ Result")
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Original**")
                    st.text_area("Original file text", value=raw_text, height=300, key="orig_file_text", label_visibility="hidden")
                with col2:
                    st.markdown("**Corrected**")
                    st.text_area("Corrected file text", value=corrected, height=300, key="corr_file_text", label_visibility="hidden")

                # Prefer client-side playback for speed
                try:
                    speak_in_browser(corrected)
                except Exception:
                    st.info("Browser TTS not available — interact with the page or use a browser with Web Speech API support.")

                docx_file = create_docx(corrected)
                txt_bytes = corrected.encode("utf-8")
                st.download_button(label="📥 Download Corrected DOCX", data=docx_file, file_name="corrected_" + uploaded_file.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.download_button(label="📄 Download Corrected TXT", data=txt_bytes, file_name="corrected_" + uploaded_file.name.rsplit('.', 1)[0] + '.txt', mime="text/plain")
# c:users\shrey\AppData\Roaming\Python\Python313\Scripts\streamlit.exe" run spelling_correction_app.py